#!/usr/bin/env python3
"""Create editable Phase 5 operator documents using the approved local DOCX fallback."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "phase5" / "documents"
OUT.mkdir(parents=True, exist_ok=True)

BLACK = RGBColor(17, 17, 17)
GOLD = RGBColor(217, 164, 65)
CHARCOAL = RGBColor(54, 54, 54)
MUTED = RGBColor(102, 102, 102)
WHITE = RGBColor(255, 255, 255)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, *, size=11, bold=False, italic=False, color=BLACK, font="Arial") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def configure(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = GOLD if style_name != "Heading 3" else CHARCOAL
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run(running_label.upper()), size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("OUR TOWN PROPERTIES, INC.  |  ASK MAGIC MIKE  |  INTERNAL OPERATIONS"), size=7.5, color=MUTED)


def add_cover(doc: Document, title: str, subtitle: str, status: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("ASK MAGIC MIKE"), size=11, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run(title), size=26, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_run(p.add_run(subtitle), size=13, color=CHARCOAL)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(2.15), Inches(2.15), Inches(2.15)]
    values = [("STATUS", status), ("DATE", "August 15, 2026"), ("DATA", "No genuine customer PII")]
    for idx, ((label, value), width) in enumerate(zip(values, widths)):
        cell = table.cell(0, idx)
        cell.width = width
        shade_cell(cell, "111111")
        set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(label + "\n"), size=7.5, bold=True, color=GOLD)
        set_run(p.add_run(value), size=9.2, bold=True, color=WHITE)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    shade_cell(cell, "F6E8BF")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label + ": "), size=10.5, bold=True, color=CHARCOAL)
    set_run(p.add_run(text), size=10.5, color=BLACK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        shade_cell(cell, "242424")
        set_cell_margins(cell)
        set_run(cell.paragraphs[0].add_run(header), size=8.5, bold=True, color=WHITE)
    for row in rows:
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row, widths)):
            cells[idx].width = Inches(width)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            if len(table.rows) % 2 == 0:
                shade_cell(cells[idx], "F8F8F8")
            set_run(cells[idx].paragraphs[0].add_run(str(value)), size=8.7, color=BLACK)
    doc.add_paragraph()


def add_markdown_body(doc: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or line.startswith("# "):
            idx += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("|") and idx + 1 < len(lines) and re.match(r"^\|[ :\-|]+\|$", lines[idx + 1].strip()):
            headers = [part.strip() for part in line.strip("|").split("|")]
            idx += 2
            rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append([part.strip().replace("`", "") for part in lines[idx].strip().strip("|").split("|")])
                idx += 1
            add_table(doc, headers, rows)
            continue
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            set_run(p.add_run(re.sub(r"^\d+\. ", "", line).replace("`", "")), size=10.5)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_run(p.add_run(line[2:].replace("`", "")), size=10.5)
        elif line.startswith("> "):
            add_callout(doc, "PROPOSED LANGUAGE", line[2:].replace("`", ""))
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            set_run(paragraph.add_run(line.replace("`", "")), size=10.5)
        idx += 1


def save_markdown_document(source_name: str, output_name: str, title: str, subtitle: str, status: str) -> Path:
    doc = Document()
    configure(doc, title)
    add_cover(doc, title, subtitle, status)
    add_markdown_body(doc, ROOT / "docs" / source_name)
    path = OUT / output_name
    doc.save(path)
    return path


def build_day1_report() -> Path:
    doc = Document()
    configure(doc, "Day 1 Operations Report")
    add_cover(doc, "Day 1 Operations Report", "Production control, queue integrity, and operator readiness", "LIVE / MONITORED")
    add_callout(doc, "EXECUTIVE RESULT", "The public funnel is reachable and monitored. The authenticated Active/New queue contains zero QA records. No genuine prospect has been fabricated or observed at this evidence point.")
    doc.add_heading("Verified production baseline", level=1)
    add_table(doc, ["Control", "Verified state", "Evidence"], [
        ["Canonical commit", "e754456cecaf6538df25bb4bf5eebe57ebf6eacb", "Git/Vercel"],
        ["Deployment", "dpl_3ogimm1EhHCaPkEfXLAeojrm2H8Z - Ready", "Vercel"],
        ["Database", "Neon production branch br-round-base-auh6h2wd", "Neon console"],
        ["Public checks", "Health passed; public monitor 9/9", "HTTP + scheduled workflow"],
        ["Lead state", "0 genuine; 6 QA; all 6 suppressed", "Privacy-safe aggregate"],
        ["Active/New", "0 QA; no STALLED or ROUTING READY QA", "Signed-in Lead Center"],
        ["Canonical forms", "Form 3 only", "WordPress bridge 1.1.0"],
        ["Web Push", "0 devices enrolled", "Production aggregate"],
    ], widths=[1.55, 3.25, 1.7])
    doc.add_heading("Operator state", level=1)
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run("Brandon: administrator, email verified, credential present, authenticated access confirmed."), size=10.5)
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run("Mike: primary lead owner provisioned; private password activation and physical Push permission remain pending."), size=10.5)
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run("Carrier SMS remains disabled. Email is the primary live notification channel; Web Push is additive after device enrollment."), size=10.5)
    doc.add_heading("First genuine lead controls", level=1)
    add_table(doc, ["Stage", "Control", "Escalation"], [
        ["Create", "Persist before notification; explicit QA evidence required for test classification", "Storage failure is critical"],
        ["Assign", "Mike or approved admin fallback with reason and audit event", "Unassigned is high severity"],
        ["Notify", "One idempotent internal email; protected audit copy; bounded retry", "Failed delivery remains visible"],
        ["Respond", "Consent-permitted human acknowledgment and outcome", "Internal SLA monitor; no public promise"],
        ["Report", "Exclude test and suppressed records from live KPIs", "Queue invariant failure is critical"],
    ], widths=[1.0, 3.65, 1.85])
    doc.add_heading("Open human approvals", level=1)
    add_table(doc, ["Owner", "Action", "Why human", "Live funnel impact"], [
        ["Mike", "Choose private Lead Center password, then enroll Push if desired", "Private credential and physical browser permission", "None; email funnel remains live"],
        ["BIC/owner", "Approve consent/routing packet before any held form activation", "Brokerage/legal decision", "None; Form 3 remains active"],
        ["Hosting operator", "Review only the Meta-blocking rule for two URLs", "Managed WAF control", "None; use AskMagicMike.com previews"],
    ], widths=[1.1, 2.4, 1.65, 1.35])
    doc.add_heading("Next operating action", level=1)
    add_callout(doc, "ONE ACTION", "Mike completes his private Lead Center password activation. No reset link or credential belongs in this report.")
    path = OUT / "ASK_MAGIC_MIKE_DAY1_OPERATIONS_REPORT.docx"
    doc.save(path)
    return path


def main() -> None:
    paths = [
        build_day1_report(),
        save_markdown_document("GRAVITY_FORMS_BIC_APPROVAL_PACKET.md", "GRAVITY_FORMS_BIC_APPROVAL_PACKET.docx", "Gravity Forms BIC Approval Packet", "Controlled form preparation and one-form-at-a-time activation decisions", "APPROVAL REQUIRED"),
        save_markdown_document("FIRST_LIVE_LEAD_RESPONSE_RUNBOOK.md", "FIRST_LIVE_LEAD_RESPONSE_RUNBOOK.docx", "First Genuine Lead Response Runbook", "Private operator response, acceptance, and escalation controls", "READY / NOT YET OBSERVED"),
        save_markdown_document("BRANDON_WEB_PUSH_ENROLLMENT.md", "BRANDON_WEB_PUSH_ENROLLMENT.docx", "Brandon Web Push Enrollment", "Physical-device enrollment guide for the administrator copy role", "PHYSICAL ACTION PENDING"),
        save_markdown_document("MIKE_WEB_PUSH_ENROLLMENT.md", "MIKE_WEB_PUSH_ENROLLMENT.docx", "Mike Web Push Enrollment", "Private credential activation followed by primary-device enrollment", "PRIVATE ACTION PENDING"),
    ]
    print("\n".join(str(path) for path in paths))


if __name__ == "__main__":
    main()
